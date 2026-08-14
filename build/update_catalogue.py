"""Brings AMS_Module_Screen_Feature_Catalogue_v2.docx up to Revision 3.

The catalogue is the business-facing document: it is what a reader who will
never open the design script uses to understand what AMS does. Revision 3
changed what the Assets module IS, and deleted a module outright, so leaving
this file alone would make it the most confidently wrong document in the repo.

Run: python build/update_catalogue.py
It is idempotent - re-running against an already-updated file changes nothing.
"""
import copy
import pathlib
import sys

import docx
from docx.text.paragraph import Paragraph

DOC =pathlib.Path(__file__).resolve().parent.parent / "docs" / "AMS_Module_Screen_Feature_Catalogue_v2.docx"

# ---------------------------------------------------------------- new content

# Screens added to module 3. (name, purpose, who)
NEW_ASSET_SCREENS = [
    ("Asset Classes & Chart of Accounts",
     "The finance taxonomy: the thirteen asset classes, the reporting category each rolls up to, and the ledger codes.",
     "Super Admin"),
    ("Bulk Stock",
     "How much of each bulk line is held, branch by branch and site by site.",
     "Super Admin, Branch Admin"),
    ("Disposals",
     "Recording a sale or write-off: quantity, proceeds, reason and who approved it.",
     "Super Admin"),
]

# Features added to module 3. (feature, who, status, what happens)
NEW_ASSET_FEATURES = [
    ("Classify an asset for the accounts", "Super Admin", "New",
     "Asset class and reporting category, kept separate from asset type because the two genuinely disagree - Storage Rack is Furniture & Fixtures and Plant & Machinery and Office Equipments. The thirteen classes are the ones the accounts already run on."),
    ("Say what a type of asset can do", "Super Admin", "New",
     "Per asset type: whether it can be issued to a person, whether it is physical, whether it is counted in bulk, and which detail record applies. Behaviour is data, so adding a kind of asset does not need a developer."),
    ("Bulk lines carry a quantity", "Super Admin, Branch Admin", "New",
     "A line of 495 barricades is one record with a quantity and a unit, not 495 records. Only bulk lines may carry a quantity above one, and the database enforces it, so allocation and verification keep working unchanged for everything else."),
    ("How much is held, and where", "Super Admin, Branch Admin", "New",
     "Each bulk line's on-hand balance at each branch and customer site. Issuing more than is held is refused by the database rather than by a check the code might skip."),
    ("Split a bulk line", "Super Admin", "New",
     "Carving individually tracked assets out of a bulk line, each keeping the link back to the line it came from."),
    ("Capitalise an asset under construction", "Super Admin", "New",
     "An asset under construction becomes a real asset and keeps the link back. Most of the register came into existence this way, so it is the normal path, not an exception."),
    ("Dispose of an asset", "Super Admin", "New",
     "Date, quantity, proceeds, reason and approver. Distinct from scrapping: a disposal was sold, carries money, and a bulk line can be disposed of in parts."),
    ("Book values, mirrored from SAP", "Super Admin", "New",
     "Gross value, accumulated depreciation and net book value shown on the asset, read-only. SAP S/4HANA owns the arithmetic; AMS never recomputes it, because two systems calculating one number is how one asset ends up with two answers."),
    ("Depreciation history", "Super Admin", "New",
     "Opening balance, charge for the year and closing balance, one row per financial year, so any prior year can be reproduced instead of inferred."),
    ("Vehicle details", "Super Admin, Branch Admin", "New",
     "Registration, chassis, engine, fuel type, fitness, PUC and insurance expiry, and odometer reading."),
    ("Field equipment lives in the same register", "Field Asset Admin", "New",
     "Site equipment is registered, searched and imported here. The field-asset capabilities scope what a field administrator sees; they no longer open a separate register."),
    ("Every imported row remembers its batch", "Super Admin", "New",
     "An imported asset keeps a link to the import that created it, so a bad file can be traced without guessing."),
]

# Existing module 3 rows that Revision 3 changed. Matched on the first cell.
ASSET_FEATURE_EDITS = {
    "Register an asset": (None, None,
                          "Asset number, name, serial, type, class, make, model, status, branch, department and cost centre."),
    "Hardware details": (None, None,
                         "Processor, memory, storage, monitor, hostname, MAC and IP. Make and model moved onto the asset record itself, because every kind of asset has them, not just computers."),
    "Calibration dates": ("Calibration and instrument details", "Super Admin, Branch Admin",
                          "Calibration window, frequency, agency, certificate number, measurement range and accuracy class. Held on a separate instrument record that only the 221 assets needing it carry."),
    "Define custom fields": (None, None,
                             "Extra fields per asset type, with type, required flag, range and dropdown options."),
    "Fill custom fields": (None, None,
                           "The asset form shows the fields defined for that asset's type."),
}

ASSET_SCREEN_EDITS = {
    "Categories & Custom Fields": ("Asset Types & Custom Fields",
                                   "The asset type tree, what each type is allowed to do, and the custom fields defined for it.",
                                   "Super Admin"),
    "Asset Register": (None,
                       "Search, filter, page and export the register - every asset the company owns, not only IT.",
                       "Super Admin, Branch Admin, Field Asset Admin"),
}

# Straight paragraph replacements elsewhere in the document.
PARAGRAPH_EDITS = [
    ("The asset register itself, and the timeline of everything that has happened to each asset.",
     "Every asset the company owns - IT hardware and software, factory and installation equipment, furniture, vehicles, instruments and leasehold property - and the timeline of everything that has happened to each one."),
    ("Loading spreadsheets: assets, employees, field assets and the fixed asset register.",
     "Loading spreadsheets: assets, employees and the fixed asset register."),
    ("Scope: The field inventory only.",
     "Scope: Field equipment within the one register."),
]

OLD_NOTE = "Two things in the source handbook were deliberately not carried over."
NEW_NOTE_PREFIX = "Three things in the source handbook were deliberately not carried over."


def cell_text(cell) -> str:
    return cell.text.strip()


def set_cell(cell, text: str):
    """Replaces a cell's text, keeping the formatting of its first run."""
    if cell_text(cell) == text:
        return False
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    return True


def append_row(table, values):
    """Clones the last row so the new one inherits its formatting."""
    template = table.rows[-1]._element
    new = copy.deepcopy(template)
    template.addnext(new)
    row = table.rows[-1]
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)
    return row


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def para_text(document, element):
    """The paragraph's text.

    NOT element.itertext(): this document nests its runs such that itertext
    yields every string three times, so a heading reads
    '16. Field Assets16. Field Assets16. Field Assets' and no equality test
    against it can ever match. python-docx walks w:r/w:t properly.
    """
    return Paragraph(element, document).text.strip()


def is_heading1(element) -> bool:
    return element.tag == W + "p" and any(
        e.get(W + "val") == "Heading1" for e in element.iter() if e.tag == W + "pStyle")


def find_module_block(document, heading_text):
    """The body elements making up one module: its Heading 1 to the next one."""
    body = list(document.element.body)
    start = next((i for i, el in enumerate(body)
                  if el.tag == W + "p" and para_text(document, el) == heading_text), None)
    if start is None:
        return []
    end = next((j for j in range(start + 1, len(body)) if is_heading1(body[j])), len(body))
    return body[start:end]


def main():
    if not DOC.exists():
        sys.exit(f"catalogue not found: {DOC}")
    d = docx.Document(DOC)
    changes = []

    # ---- 1. delete module 16, Field Assets ------------------------------
    block = find_module_block(d, "16. Field Assets")
    if block:
        for el in block:
            el.getparent().remove(el)
        changes.append(f"deleted module 16 Field Assets ({len(block)} body elements)")
    elif any("Field Assets" in para_text(d, el) for el in d.element.body if is_heading1(el)):
        # Silence here would ship a catalogue that still documents a module
        # that no longer exists, and the counts below would quietly agree
        # with it. Fail instead.
        sys.exit("found a Field Assets heading but could not match it - the heading text has changed")

    # ---- 2. renumber Reports and Dashboards -----------------------------
    for para in d.paragraphs:
        if para.text.strip() == "17. Reports and Dashboards":
            for run in para.runs:
                if "17." in run.text:
                    run.text = run.text.replace("17.", "16.")
                    break
            changes.append("renumbered Reports and Dashboards 17 -> 16")
            break

    # ---- 3. module 3 screens and features -------------------------------
    screens = d.tables[6]
    features = d.tables[7]

    for row in screens.rows[1:]:
        edit = ASSET_SCREEN_EDITS.get(cell_text(row.cells[0]))
        if edit:
            name, purpose, who = edit
            if name:
                set_cell(row.cells[0], name)
            set_cell(row.cells[1], purpose)
            set_cell(row.cells[2], who)
            changes.append(f"screen updated: {name or cell_text(row.cells[0])}")

    existing_screens = {cell_text(r.cells[0]) for r in screens.rows}
    for name, purpose, who in NEW_ASSET_SCREENS:
        if name not in existing_screens:
            append_row(screens, [name, purpose, who])
            changes.append(f"screen added: {name}")

    for row in features.rows[1:]:
        edit = ASSET_FEATURE_EDITS.get(cell_text(row.cells[0]))
        if edit:
            name, who, what = edit
            if name:
                set_cell(row.cells[0], name)
            if who:
                set_cell(row.cells[1], who)
            set_cell(row.cells[3], what)
            changes.append(f"feature updated: {name or cell_text(row.cells[0])}")

    existing_features = {cell_text(r.cells[0]) for r in features.rows}
    for name, who, status, what in NEW_ASSET_FEATURES:
        if name not in existing_features:
            append_row(features, [name, who, status, what])
            changes.append(f"feature added: {name}")

    # ---- 4. paragraph edits ---------------------------------------------
    for old, new in PARAGRAPH_EDITS:
        for para in d.paragraphs:
            if para.text.strip() == old:
                if para.runs:
                    para.runs[0].text = new
                    for run in para.runs[1:]:
                        run.text = ""
                changes.append(f"paragraph: {old[:45]}...")
                break

    # Revision 2 -> Revision 3, everywhere it refers to the database design.
    for para in d.paragraphs:
        if "Revision 2" in para.text:
            for run in para.runs:
                if "Revision 2" in run.text:
                    run.text = run.text.replace("Revision 2", "Revision 3")
            changes.append("Revision 2 -> Revision 3 in a paragraph")
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if "Revision 2" in cell.text:
                    set_cell(cell, cell.text.strip().replace("Revision 2", "Revision 3"))
                    changes.append("Revision 2 -> Revision 3 in a table cell")

    # The "deliberately not carried over" note gains the second register.
    for para in d.paragraphs:
        if para.text.strip().startswith(OLD_NOTE):
            text = para.text.strip().replace(OLD_NOTE, NEW_NOTE_PREFIX)
            text += (" And field assets no longer get a register of their own either: "
                     "a second asset register is the same mistake as a second login table, "
                     "one level up. Field equipment is registered alongside everything else, "
                     "and the field-asset capabilities scope what a field administrator sees.")
            if para.runs:
                para.runs[0].text = text
                for run in para.runs[1:]:
                    run.text = ""
            changes.append("scope note now names three omissions")
            break

    # ---- 5. counts, computed from the document itself --------------------
    # Tables alternate screens, features per module after the two intro tables.
    module_tables = d.tables[2:]
    screen_count = sum(len(t.rows) - 1 for t in module_tables[0::2])
    feature_rows = [r for t in module_tables[1::2] for r in t.rows[1:]]
    feature_count = len(feature_rows)
    live = sum(1 for r in feature_rows if cell_text(r.cells[2]) == "Live")
    new = feature_count - live
    module_count = len(module_tables) // 2

    for para in d.paragraphs:
        if para.text.strip().startswith("Counted across the whole catalogue"):
            text = (f"Counted across the whole catalogue: {feature_count} features and "
                    f"{screen_count} screens in {module_count} modules. {live} features are "
                    f"live today and {new} are new in this design.")
            if para.runs:
                para.runs[0].text = text
                for run in para.runs[1:]:
                    run.text = ""
            changes.append(f"counts: {feature_count} features, {screen_count} screens, "
                           f"{module_count} modules, {live} live, {new} new")
            break

    d.save(DOC)
    for c in changes:
        print(f"  {c}")
    print(f"\n{len(changes)} changes written to {DOC.name}")


if __name__ == "__main__":
    main()
